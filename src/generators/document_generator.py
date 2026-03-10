"""
Генератор документов в форматах DOCX и PDF.
Поддержка ГОСТ 7.32-2017 для оформления отчётов.
"""

from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
import datetime
import io

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    DocxDocument = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, 
        TableStyle, PageBreak, Image
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    SimpleDocTemplate = None


class DocumentFormat(Enum):
    DOCX = "docx"
    PDF = "pdf"


@dataclass
class DocumentSection:
    """Секция документа."""
    title: str
    content: str
    level: int = 1  # Уровень заголовка (1-3)
    subsections: List['DocumentSection'] = field(default_factory=list)


@dataclass
class DocumentMetadata:
    """Метаданные документа."""
    title: str
    author: str
    organization: Optional[str] = None
    date: Optional[datetime.date] = None
    version: str = "1.0"
    keywords: List[str] = field(default_factory=list)
    
    def __post_init__(self):
        if self.date is None:
            self.date = datetime.date.today()


@dataclass
class GOSTSettings:
    """
    Настройки оформления по ГОСТ 7.32-2017.
    """
    # Поля страницы (в см)
    margin_top: float = 2.0
    margin_bottom: float = 2.0
    margin_left: float = 3.0  # Для подшивки
    margin_right: float = 1.5
    
    # Шрифт
    font_name: str = "Times New Roman"
    font_size: int = 14  # Основной текст
    font_size_heading1: int = 16
    font_size_heading2: int = 14
    font_size_heading3: int = 14
    
    # Интервалы
    line_spacing: float = 1.5
    paragraph_spacing_before: int = 0
    paragraph_spacing_after: int = 0
    first_line_indent: float = 1.25  # Абзацный отступ в см
    
    # Нумерация
    page_numbering: bool = True
    page_number_position: str = "bottom_center"


class GOSTDocxGenerator:
    """
    Генератор DOCX документов по ГОСТ 7.32-2017.
    """
    
    def __init__(self, settings: Optional[GOSTSettings] = None):
        if DocxDocument is None:
            raise ImportError(
                "Установите python-docx: pip install python-docx"
            )
        
        self.settings = settings or GOSTSettings()
        self.document = None
    
    def create_document(self, metadata: DocumentMetadata) -> DocxDocument:
        """Создание нового документа с настройками ГОСТ."""
        self.document = DocxDocument()
        self._setup_styles()
        self._setup_page_layout()
        self._add_title_page(metadata)
        return self.document
    
    def _setup_styles(self):
        """Настройка стилей документа."""
        styles = self.document.styles
        
        # Стиль обычного текста
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.settings.font_name
        normal_font.size = Pt(self.settings.font_size)
        
        paragraph_format = normal_style.paragraph_format
        paragraph_format.line_spacing = self.settings.line_spacing
        paragraph_format.first_line_indent = Cm(self.settings.first_line_indent)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Стили заголовков
        self._create_heading_style('GOST Heading 1', 1, self.settings.font_size_heading1)
        self._create_heading_style('GOST Heading 2', 2, self.settings.font_size_heading2)
        self._create_heading_style('GOST Heading 3', 3, self.settings.font_size_heading3)
    
    def _create_heading_style(self, name: str, level: int, font_size: int):
        """Создание стиля заголовка."""
        styles = self.document.styles
        
        try:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style = styles[name]
        
        font = style.font
        font.name = self.settings.font_name
        font.size = Pt(font_size)
        font.bold = True
        
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(12 if level == 1 else 6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.first_line_indent = Cm(0)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if level == 1:
            paragraph_format.page_break_before = True
    
    def _setup_page_layout(self):
        """Настройка параметров страницы."""
        for section in self.document.sections:
            section.top_margin = Cm(self.settings.margin_top)
            section.bottom_margin = Cm(self.settings.margin_bottom)
            section.left_margin = Cm(self.settings.margin_left)
            section.right_margin = Cm(self.settings.margin_right)
    
    def _add_title_page(self, metadata: DocumentMetadata):
        """Добавление титульной страницы."""
        # Организация
        if metadata.organization:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata.organization.upper())
            run.font.size = Pt(14)
            run.font.bold = True
        
        # Отступ сверху
        for _ in range(8):
            self.document.add_paragraph()
        
        # Название документа
        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(metadata.title.upper())
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        
        # Отступ
        for _ in range(4):
            self.document.add_paragraph()
        
        # Информация об авторе
        author_para = self.document.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        author_para.add_run(f"Выполнил: {metadata.author}")
        
        # Дата
        for _ in range(6):
            self.document.add_paragraph()
        
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(metadata.date.strftime("%Y"))
        
        # Разрыв страницы
        self.document.add_page_break()
    
    def add_section(self, section: DocumentSection, numbering: str = ""):
        """Добавление секции в документ."""
        # Формируем номер секции
        section_number = numbering if numbering else ""
        
        # Добавляем заголовок
        heading_style = f'GOST Heading {section.level}'
        heading_text = f"{section_number} {section.title}".strip()
        
        heading = self.document.add_paragraph(heading_text, style=heading_style)
        
        # Добавляем содержимое
        if section.content:
            paragraphs = section.content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = self.document.add_paragraph(para_text.strip())
                    p.style = 'Normal'
        
        # Рекурсивно добавляем подсекции
        for i, subsection in enumerate(section.subsections, 1):
            sub_number = f"{section_number}.{i}" if section_number else str(i)
            self.add_section(subsection, sub_number)
    
    def add_table(
        self,
        data: List[List[str]],
        headers: Optional[List[str]] = None,
        caption: Optional[str] = None
    ):
        """Добавление таблицы."""
        if caption:
            caption_para = self.document.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.add_run(f"Таблица — {caption}")
        
        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else len(headers) if headers else 0
        
        if cols == 0:
            return
        
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        row_idx = 0
        
        # Заголовки
        if headers:
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                # Жирный текст в заголовке
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            row_idx = 1
        
        # Данные
        for data_row in data:
            for col_idx, cell_text in enumerate(data_row):
                table.rows[row_idx].cells[col_idx].text = str(cell_text)
            row_idx += 1
        
        # Отступ после таблицы
        self.document.add_paragraph()
    
    def add_code_block(self, code: str, language: str = ""):
        """Добавление блока кода."""
        # Добавляем рамку для кода
        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        
        run = para.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    
    def add_list(self, items: List[str], ordered: bool = False):
        """Добавление списка."""
        for i, item in enumerate(items):
            if ordered:
                prefix = f"{i + 1}. "
            else:
                prefix = "— "
            
            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            para.add_run(f"{prefix}{item}")
    
    def save(self, filepath: Union[str, Path]) -> Path:
        """Сохранение документа."""
        filepath = Path(filepath)
        if not filepath.suffix:
            filepath = filepath.with_suffix('.docx')
        
        self.document.save(str(filepath))
        return filepath
    
    def get_bytes(self) -> bytes:
        """Получение документа в виде байтов."""
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.read()


class PDFGenerator:
    """
    Генератор PDF документов.
    """
    
    def __init__(self, settings: Optional[GOSTSettings] = None):
        if SimpleDocTemplate is None:
            raise ImportError(
                "Установите reportlab: pip install reportlab"
            )
        
        self.settings = settings or GOSTSettings()
        self.elements = []
        self.styles = self._create_styles()
    
    def _create_styles(self) -> Dict[str, ParagraphStyle]:
        """Создание стилей для PDF."""
        base_styles = getSampleStyleSheet()
        
        styles = {
            'Normal': ParagraphStyle(
                'CustomNormal',
                parent=base_styles['Normal'],
                fontSize=self.settings.font_size,
                leading=self.settings.font_size * self.settings.line_spacing,
                firstLineIndent=self.settings.first_line_indent * cm,
                alignment=4  # JUSTIFY
            ),
            'Heading1': ParagraphStyle(
                'CustomHeading1',
                parent=base_styles['Heading1'],
                fontSize=self.settings.font_size_heading1,
                spaceAfter=12,
                spaceBefore=24
            ),
            'Heading2': ParagraphStyle(
                'CustomHeading2',
                parent=base_styles['Heading2'],
                fontSize=self.settings.font_size_heading2,
                spaceAfter=6,
                spaceBefore=12
            ),
            'Code': ParagraphStyle(
                'Code',
                parent=base_styles['Code'],
                fontSize=10,
                leftIndent=1 * cm
            ),
            'Title': ParagraphStyle(
                'CustomTitle',
                parent=base_styles['Title'],
                fontSize=18,
                alignment=1  # CENTER
            )
        }
        
        return styles
    
    def add_title_page(self, metadata: DocumentMetadata):
        """Добавление титульной страницы."""
        # Организация
        if metadata.organization:
            self.elements.append(
                Paragraph(metadata.organization.upper(), self.styles['Normal'])
            )
        
        self.elements.append(Spacer(1, 5 * cm))
        
        # Заголовок
        self.elements.append(
            Paragraph(metadata.title.upper(), self.styles['Title'])
        )
        
        self.elements.append(Spacer(1, 3 * cm))
        
        # Автор
        author_style = ParagraphStyle(
            'Author',
            parent=self.styles['Normal'],
            alignment=2  # RIGHT
        )
        self.elements.append(
            Paragraph(f"Выполнил: {metadata.author}", author_style)
        )
        
        self.elements.append(Spacer(1, 5 * cm))
        
        # Дата
        date_style = ParagraphStyle(
            'Date',
            parent=self.styles['Normal'],
            alignment=1  # CENTER
        )
        self.elements.append(
            Paragraph(metadata.date.strftime("%Y"), date_style)
        )
        
        self.elements.append(PageBreak())
    
    def add_section(self, section: DocumentSection, numbering: str = ""):
        """Добавление секции."""
        # Заголовок
        style_name = f'Heading{min(section.level, 2)}'
        heading_text = f"{numbering} {section.title}".strip()
        self.elements.append(
            Paragraph(heading_text, self.styles[style_name])
        )
        
        # Содержимое
        if section.content:
            paragraphs = section.content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    self.elements.append(
                        Paragraph(para_text.strip(), self.styles['Normal'])
                    )
                    self.elements.append(Spacer(1, 0.3 * cm))
        
        # Подсекции
        for i, subsection in enumerate(section.subsections, 1):
            sub_number = f"{numbering}.{i}" if numbering else str(i)
            self.add_section(subsection, sub_number)
    
    def add_table(
        self,
        data: List[List[str]],
        headers: Optional[List[str]] = None,
        caption: Optional[str] = None
    ):
        """Добавление таблицы."""
        if caption:
            caption_style = ParagraphStyle(
                'Caption',
                parent=self.styles['Normal'],
                alignment=1
            )
            self.elements.append(
                Paragraph(f"Таблица — {caption}", caption_style)
            )
            self.elements.append(Spacer(1, 0.3 * cm))
        
        table_data = []
        
        if headers:
            table_data.append(headers)
        
        table_data.extend(data)
        
        table = Table(table_data)
        
        style = TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('PADDING', (0, 0), (-1, -1), 6),
        ])
        
        if headers:
            style.add('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey)
        
        table.setStyle(style)
        self.elements.append(table)
        self.elements.append(Spacer(1, 0.5 * cm))
    
    def save(self, filepath: Union[str, Path]) -> Path:
        """Сохранение PDF."""
        filepath = Path(filepath)
        if not filepath.suffix:
            filepath = filepath.with_suffix('.pdf')
        
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            topMargin=self.settings.margin_top * cm,
            bottomMargin=self.settings.margin_bottom * cm,
            leftMargin=self.settings.margin_left * cm,
            rightMargin=self.settings.margin_right * cm
        )
        
        doc.build(self.elements)
        return filepath
    
    def get_bytes(self) -> bytes:
        """Получение PDF в виде байтов."""
        buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=self.settings.margin_top * cm,
            bottomMargin=self.settings.margin_bottom * cm,
            leftMargin=self.settings.margin_left * cm,
            rightMargin=self.settings.margin_right * cm
        )
        
        doc.build(self.elements)
        buffer.seek(0)
        return buffer.read()


class DocumentGenerator:
    """
    Унифицированный генератор документов.
    Поддерживает DOCX и PDF форматы.
    """
    
    def __init__(
        self,
        settings: Optional[GOSTSettings] = None,
        llm_service: Any = None  # Для AI-генерации контента
    ):
        self.settings = settings or GOSTSettings()
        self.llm_service = llm_service
    
    def generate(
        self,
        metadata: DocumentMetadata,
        sections: List[DocumentSection],
        format: DocumentFormat = DocumentFormat.DOCX,
        output_path: Optional[Union[str, Path]] = None
    ) -> Union[Path, bytes]:
        """
        Генерация документа.
        
        Args:
            metadata: Метаданные документа
            sections: Список секций
            format: Формат документа (DOCX или PDF)
            output_path: Путь для сохранения (опционально)
            
        Returns:
            Путь к файлу или байты документа
        """
        if format == DocumentFormat.DOCX:
            generator = GOSTDocxGenerator(self.settings)
            generator.create_document(metadata)
            
            for i, section in enumerate(sections, 1):
                generator.add_section(section, str(i))
            
            if output_path:
                return generator.save(output_path)
            return generator.get_bytes()
        
        else:  # PDF
            generator = PDFGenerator(self.settings)
            generator.add_title_page(metadata)
            
            for i, section in enumerate(sections, 1):
                generator.add_section(section, str(i))
            
            if output_path:
                return generator.save(output_path)
            return generator.get_bytes()
    
    def generate_from_template(
        self,
        template_name: str,
        context: Dict[str, Any],
        format: DocumentFormat = DocumentFormat.DOCX
    ) -> bytes:
        """
        Генерация документа из шаблона с AI-заполнением.
        """
        templates = {
            "srs": self._create_srs_template,
            "test_plan": self._create_test_plan_template,
            "technical_doc": self._create_technical_doc_template,
            "release_notes": self._create_release_notes_template,
        }
        
        if template_name not in templates:
            raise ValueError(f"Неизвестный шаблон: {template_name}")
        
        metadata, sections = templates[template_name](context)
        
        return self.generate(metadata, sections, format)
    
    def _create_srs_template(
        self,
        context: Dict[str, Any]
    ) -> tuple[DocumentMetadata, List[DocumentSection]]:
        """Шаблон SRS (Software Requirements Specification)."""
        metadata = DocumentMetadata(
            title=f"Спецификация требований к ПО: {context.get('project_name', 'Проект')}",
            author=context.get('author', 'Автор'),
            organization=context.get('organization'),
            version=context.get('version', '1.0')
        )
        
        sections = [
            DocumentSection(
                title="Введение",
                content=context.get('introduction', ''),
                level=1,
                subsections=[
                    DocumentSection(
                        title="Цель документа",
                        content="Данный документ описывает требования к программному обеспечению...",
                        level=2
                    ),
                    DocumentSection(
                        title="Область применения",
                        content=context.get('scope', ''),
                        level=2
                    )
                ]
            ),
            DocumentSection(
                title="Функциональные требования",
                content="",
                level=1,
                subsections=[
                    DocumentSection(
                        title=f"FR-{i+1}: {req.get('name', '')}",
                        content=req.get('description', ''),
                        level=2
                    )
                    for i, req in enumerate(context.get('functional_requirements', []))
                ]
            ),
            DocumentSection(
                title="Нефункциональные требования",
                content="",
                level=1,
                subsections=[
                    DocumentSection(
                        title=f"NFR-{i+1}: {req.get('name', '')}",
                        content=req.get('description', ''),
                        level=2
                    )
                    for i, req in enumerate(context.get('non_functional_requirements', []))
                ]
            )
        ]
        
        return metadata, sections
    
    def _create_technical_doc_template(
        self,
        context: Dict[str, Any]
    ) -> tuple[DocumentMetadata, List[DocumentSection]]:
        """Шаблон технической документации."""
        metadata = DocumentMetadata(
            title=f"Техническая документация: {context.get('component_name', 'Компонент')}",
            author=context.get('author', 'Автор'),
            organization=context.get('organization')
        )
        
        sections = [
            DocumentSection(
                title="Обзор",
                content=context.get('overview', ''),
                level=1
            ),
            DocumentSection(
                title="Архитектура",
                content=context.get('architecture', ''),
                level=1
            ),
            DocumentSection(
                title="API Reference",
                content=context.get('api_reference', ''),
                level=1
            ),
            DocumentSection(
                title="Примеры использования",
                content=context.get('examples', ''),
                level=1
            )
        ]
        
        return metadata, sections
    
    def _create_test_plan_template(
        self,
        context: Dict[str, Any]
    ) -> tuple[DocumentMetadata, List[DocumentSection]]:
        """Шаблон тест-плана."""
        metadata = DocumentMetadata(
            title=f"План тестирования: {context.get('project_name', 'Проект')}",
            author=context.get('author', 'Автор'),
            organization=context.get('organization')
        )
        
        sections = [
            DocumentSection(
                title="Введение",
                content=context.get('introduction', ''),
                level=1
            ),
            DocumentSection(
                title="Объект тестирования",
                content=context.get('test_object', ''),
                level=1
            ),
            DocumentSection(
                title="Стратегия тестирования",
                content=context.get('strategy', ''),
                level=1
            ),
            DocumentSection(
                title="Тест-кейсы",
                content="",
                level=1,
                subsections=[
                    DocumentSection(
                        title=f"TC-{tc.get('id', i+1)}: {tc.get('name', '')}",
                        content=f"Предусловия: {tc.get('preconditions', '')}\n\n"
                                f"Шаги: {tc.get('steps', '')}\n\n"
                                f"Ожидаемый результат: {tc.get('expected', '')}",
                        level=2
                    )
                    for i, tc in enumerate(context.get('test_cases', []))
                ]
            )
        ]
        
        return metadata, sections
    
    def _create_release_notes_template(
        self,
        context: Dict[str, Any]
    ) -> tuple[DocumentMetadata, List[DocumentSection]]:
        """Шаблон release notes."""
        metadata = DocumentMetadata(
            title=f"Release Notes - Версия {context.get('version', '1.0')}",
            author=context.get('author', 'Команда разработки'),
            organization=context.get('organization')
        )
        
        sections = [
            DocumentSection(
                title="Обзор версии",
                content=context.get('overview', ''),
                level=1
            ),
            DocumentSection(
                title="Новые функции",
                content='\n\n'.join([
                    f"• {feature}" 
                    for feature in context.get('features', [])
                ]),
                level=1
            ),
            DocumentSection(
                title="Исправленные ошибки",
                content='\n\n'.join([
                    f"• {fix}" 
                    for fix in context.get('bug_fixes', [])
                ]),
                level=1
            ),
            DocumentSection(
                title="Известные проблемы",
                content='\n\n'.join([
                    f"• {issue}" 
                    for issue in context.get('known_issues', [])
                ]),
                level=1
            )
        ]
        
        return metadata, sections